"""
报告导出工具 - 支持 Markdown、Word、PDF 格式

依赖安装:
    pip install pypandoc markdown

PDF 导出需要额外工具:
    - wkhtmltopdf (推荐): https://wkhtmltopdf.org/downloads.html
    - 或 LaTeX: https://www.latex-project.org/get/
"""

import logging
import os
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

# 检查依赖是否可用
try:
    import markdown
    import pypandoc

    # 检查 pandoc 是否可用
    try:
        pypandoc.get_pandoc_version()
        PANDOC_AVAILABLE = True
        logger.info("✅ Pandoc 可用")
    except OSError:
        PANDOC_AVAILABLE = False
        logger.warning("⚠️ Pandoc 不可用，Word 和 PDF 导出功能将不可用")

    EXPORT_AVAILABLE = True
except ImportError as e:
    EXPORT_AVAILABLE = False
    PANDOC_AVAILABLE = False
    logger.warning(f"⚠️ 导出功能依赖包缺失: {e}")
    logger.info("💡 请安装: pip install pypandoc markdown")

# 检查替代的 PDF 生成工具
WEASYPRINT_AVAILABLE = False
PDFKIT_AVAILABLE = False
WEASYPRINT_ERROR = None
PDFKIT_ERROR = None

try:
    import weasyprint
    # 尝试创建一个简单的 HTML 来测试 Cairo 库
    try:
        weasyprint.HTML(string="<html><body>test</body></html>").write_pdf()
        WEASYPRINT_AVAILABLE = True
        logger.info("✅ WeasyPrint 可用（推荐的 PDF 生成工具）")
    except OSError as e:
        WEASYPRINT_ERROR = str(e)
        if "cairo" in str(e).lower():
            logger.warning("⚠️ WeasyPrint 已安装但缺少 Cairo 库（Windows 需要 GTK3 运行时）")
            logger.warning("   下载地址: https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer/releases")
        else:
            logger.warning(f"⚠️ WeasyPrint 不可用: {e}")
except ImportError:
    logger.info("ℹ️ WeasyPrint 不可用，可选安装: pip install weasyprint")
except Exception as e:
    WEASYPRINT_ERROR = str(e)
    logger.warning(f"⚠️ WeasyPrint 检测失败: {e}")

try:
    import pdfkit
    # 检查 wkhtmltopdf 是否安装
    try:
        pdfkit.configuration()
        PDFKIT_AVAILABLE = True
        logger.info("✅ pdfkit + wkhtmltopdf 可用")
    except Exception as e:
        PDFKIT_ERROR = str(e)
        logger.info("ℹ️ wkhtmltopdf 未安装")
except ImportError:
    logger.info("ℹ️ pdfkit 不可用，可选安装: pip install pdfkit")
except Exception as e:
    PDFKIT_ERROR = str(e)
    logger.warning(f"⚠️ pdfkit 检测失败: {e}")


class ReportExporter:
    """报告导出器 - 支持 Markdown、Word、PDF 格式"""

    def __init__(self):
        self.export_available = EXPORT_AVAILABLE
        self.pandoc_available = PANDOC_AVAILABLE
        self.weasyprint_available = WEASYPRINT_AVAILABLE
        self.pdfkit_available = PDFKIT_AVAILABLE

        logger.info(f"📋 ReportExporter 初始化:")
        logger.info(f"  - export_available: {self.export_available}")
        logger.info(f"  - pandoc_available: {self.pandoc_available}")
        logger.info(f"  - weasyprint_available: {self.weasyprint_available}")
        logger.info(f"  - pdfkit_available: {self.pdfkit_available}")
    
    def generate_markdown_report(self, report_doc: Dict[str, Any]) -> str:
        """生成 Markdown 格式报告"""
        logger.info("📝 生成 Markdown 报告...")
        
        stock_symbol = report_doc.get("stock_symbol", "unknown")
        analysis_date = report_doc.get("analysis_date", "")
        analysts = report_doc.get("analysts", [])
        research_depth = report_doc.get("research_depth", 1)
        reports = report_doc.get("reports", {})
        summary = report_doc.get("summary", "")
        
        content_parts = []
        
        # 标题和元信息
        content_parts.append(f"# {stock_symbol} 股票分析报告")
        content_parts.append("")
        content_parts.append(f"**分析日期**: {analysis_date}")
        if analysts:
            content_parts.append(f"**分析师**: {', '.join(analysts)}")
        content_parts.append(f"**研究深度**: {research_depth}")
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        
        # 执行摘要
        if summary:
            content_parts.append("## 📊 执行摘要")
            content_parts.append("")
            content_parts.append(summary)
            content_parts.append("")
            content_parts.append("---")
            content_parts.append("")
        
        # 各模块内容
        module_order = [
            "company_overview",
            "financial_analysis", 
            "technical_analysis",
            "market_analysis",
            "risk_analysis",
            "valuation_analysis",
            "investment_recommendation"
        ]
        
        module_titles = {
            "company_overview": "🏢 公司概况",
            "financial_analysis": "💰 财务分析",
            "technical_analysis": "📈 技术分析",
            "market_analysis": "🌍 市场分析",
            "risk_analysis": "⚠️ 风险分析",
            "valuation_analysis": "💎 估值分析",
            "investment_recommendation": "🎯 投资建议"
        }
        
        # 按顺序添加模块
        for module_key in module_order:
            if module_key in reports:
                module_content = reports[module_key]
                if isinstance(module_content, str) and module_content.strip():
                    title = module_titles.get(module_key, module_key)
                    content_parts.append(f"## {title}")
                    content_parts.append("")
                    content_parts.append(module_content)
                    content_parts.append("")
                    content_parts.append("---")
                    content_parts.append("")
        
        # 添加其他未列出的模块
        for module_key, module_content in reports.items():
            if module_key not in module_order:
                if isinstance(module_content, str) and module_content.strip():
                    content_parts.append(f"## {module_key}")
                    content_parts.append("")
                    content_parts.append(module_content)
                    content_parts.append("")
                    content_parts.append("---")
                    content_parts.append("")
        
        # 页脚
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        content_parts.append("*本报告由 TradingAgents-CN 自动生成*")
        content_parts.append("")
        
        markdown_content = "\n".join(content_parts)
        logger.info(f"✅ Markdown 报告生成完成，长度: {len(markdown_content)} 字符")
        
        return markdown_content
    
    def _clean_markdown_for_pandoc(self, md_content: str) -> str:
        """清理 Markdown 内容，避免 pandoc 解析问题"""
        import re

        # 移除可能导致 YAML 解析问题的内容
        # 如果开头有 "---"，在前面添加空行
        if md_content.strip().startswith("---"):
            md_content = "\n" + md_content

        # 🔥 移除可能导致竖排的 HTML 标签和样式
        # 移除 writing-mode 相关的样式
        md_content = re.sub(r'<[^>]*writing-mode[^>]*>', '', md_content, flags=re.IGNORECASE)
        md_content = re.sub(r'<[^>]*text-orientation[^>]*>', '', md_content, flags=re.IGNORECASE)

        # 移除 <div> 标签中的 style 属性（可能包含竖排样式）
        md_content = re.sub(r'<div\s+style="[^"]*">', '<div>', md_content, flags=re.IGNORECASE)
        md_content = re.sub(r'<span\s+style="[^"]*">', '<span>', md_content, flags=re.IGNORECASE)

        # 🔥 移除可能导致问题的 HTML 标签
        # 保留基本的 Markdown 格式，移除复杂的 HTML
        md_content = re.sub(r'<style[^>]*>.*?</style>', '', md_content, flags=re.DOTALL | re.IGNORECASE)

        # 🔥 确保所有段落都是正常的横排文本
        # 在每个段落前后添加明确的换行，避免 Pandoc 误判
        lines = md_content.split('\n')
        cleaned_lines = []
        for line in lines:
            # 跳过空行
            if not line.strip():
                cleaned_lines.append(line)
                continue

            # 如果是标题、列表、表格等 Markdown 语法，保持原样
            if line.strip().startswith(('#', '-', '*', '|', '>', '```', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                cleaned_lines.append(line)
            else:
                # 普通段落：确保没有特殊字符导致竖排
                cleaned_lines.append(line)

        md_content = '\n'.join(cleaned_lines)

        return md_content

    def _create_pdf_css(self) -> str:
        """创建 PDF 样式表，控制表格分页和文本方向"""
        return """
<style>
/* 🔥 强制所有文本横排显示（修复中文竖排问题） */
* {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

body {
    writing-mode: horizontal-tb !important;
    direction: ltr !important;
}

/* 段落和文本 */
p, div, span, td, th, li {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

/* 表格样式 - 允许跨页 */
table {
    width: 100%;
    border-collapse: collapse;
    page-break-inside: auto;
    writing-mode: horizontal-tb !important;
}

/* 表格行 - 避免在行中间分页 */
tr {
    page-break-inside: avoid;
    page-break-after: auto;
}

/* 表头 - 在每页重复显示 */
thead {
    display: table-header-group;
}

/* 表格单元格 */
td, th {
    padding: 8px;
    border: 1px solid #ddd;
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

/* 表头样式 */
th {
    background-color: #f2f2f2;
    font-weight: bold;
}

/* 避免标题后立即分页 */
h1, h2, h3, h4, h5, h6 {
    page-break-after: avoid;
    writing-mode: horizontal-tb !important;
}

/* 避免在列表项中间分页 */
li {
    page-break-inside: avoid;
}

/* 代码块 */
pre, code {
    writing-mode: horizontal-tb !important;
    white-space: pre-wrap;
    word-wrap: break-word;
}
</style>
"""
    
    def generate_docx_report(self, report_doc: Dict[str, Any]) -> bytes:
        """生成 Word 文档格式报告"""
        logger.info("📄 开始生成 Word 文档...")

        if not self.pandoc_available:
            raise Exception("Pandoc 不可用，无法生成 Word 文档。请安装 pandoc 或使用 Markdown 格式导出。")

        # 生成 Markdown 内容
        md_content = self.generate_markdown_report(report_doc)

        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                output_file = tmp_file.name

            logger.info(f"📁 临时文件路径: {output_file}")

            # Pandoc 参数
            extra_args = [
                '--from=markdown-yaml_metadata_block',  # 禁用 YAML 元数据块解析
                '--standalone',  # 生成独立文档
                '--wrap=preserve',  # 保留换行
                '--columns=120',  # 设置列宽
                '-M', 'lang=zh-CN',  # 🔥 明确指定语言为简体中文
                '-M', 'dir=ltr',  # 🔥 明确指定文本方向为从左到右
            ]

            # 清理内容
            cleaned_content = self._clean_markdown_for_pandoc(md_content)

            # 转换为 Word
            pypandoc.convert_text(
                cleaned_content,
                'docx',
                format='markdown',
                outputfile=output_file,
                extra_args=extra_args
            )

            logger.info("✅ pypandoc 转换完成")

            # 🔥 后处理：修复 Word 文档中的文本方向
            try:
                from docx import Document
                doc = Document(output_file)

                # 修复所有段落的文本方向
                for paragraph in doc.paragraphs:
                    # 设置段落为从左到右
                    if paragraph._element.pPr is not None:
                        # 移除可能的竖排设置
                        for child in list(paragraph._element.pPr):
                            if 'textDirection' in child.tag or 'bidi' in child.tag:
                                paragraph._element.pPr.remove(child)

                # 修复表格中的文本方向
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph._element.pPr is not None:
                                    for child in list(paragraph._element.pPr):
                                        if 'textDirection' in child.tag or 'bidi' in child.tag:
                                            paragraph._element.pPr.remove(child)

                # 保存修复后的文档
                doc.save(output_file)
                logger.info("✅ Word 文档文本方向修复完成")
            except ImportError:
                logger.warning("⚠️ python-docx 未安装，跳过文本方向修复")
            except Exception as e:
                logger.warning(f"⚠️ Word 文档文本方向修复失败: {e}")

            # 读取生成的文件
            with open(output_file, 'rb') as f:
                docx_content = f.read()

            logger.info(f"✅ Word 文档生成成功，大小: {len(docx_content)} 字节")

            # 清理临时文件
            os.unlink(output_file)

            return docx_content
            
        except Exception as e:
            logger.error(f"❌ Word 文档生成失败: {e}", exc_info=True)
            # 清理临时文件
            try:
                if 'output_file' in locals() and os.path.exists(output_file):
                    os.unlink(output_file)
            except:
                pass
            raise Exception(f"生成 Word 文档失败: {e}")
    
    def _markdown_to_html(self, md_content: str) -> str:
        """将 Markdown 转换为 HTML"""
        import markdown

        # 配置 Markdown 扩展
        extensions = [
            'markdown.extensions.tables',  # 表格支持
            'markdown.extensions.fenced_code',  # 代码块支持
            'markdown.extensions.nl2br',  # 换行支持
        ]

        # 转换为 HTML
        html_content = markdown.markdown(md_content, extensions=extensions)

        # 添加 HTML 模板和样式
        html_template = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>分析报告</title>
    <style>
        /* 🔥 强制横排显示 */
        * {{
            writing-mode: horizontal-tb !important;
            text-orientation: mixed !important;
            direction: ltr !important;
        }}

        body {{
            font-family: "Microsoft YaHei", "SimHei", "Arial", sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background: white;
        }}

        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 24px;
            margin-bottom: 16px;
            font-weight: 600;
            page-break-after: avoid;
        }}

        h1 {{ font-size: 2em; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ font-size: 1.5em; border-bottom: 1px solid #bdc3c7; padding-bottom: 8px; }}
        h3 {{ font-size: 1.25em; }}

        p {{
            margin: 12px 0;
            text-align: justify;
        }}

        /* 表格样式 */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            page-break-inside: auto;
        }}

        tr {{
            page-break-inside: avoid;
            page-break-after: auto;
        }}

        thead {{
            display: table-header-group;
        }}

        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}

        th {{
            background-color: #3498db;
            color: white;
            font-weight: bold;
        }}

        tr:nth-child(even) {{
            background-color: #f2f2f2;
        }}

        /* 代码块样式 */
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: "Consolas", "Monaco", monospace;
        }}

        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
            page-break-inside: avoid;
        }}

        pre code {{
            background-color: transparent;
            padding: 0;
        }}

        /* 列表样式 */
        ul, ol {{
            margin: 12px 0;
            padding-left: 30px;
        }}

        li {{
            margin: 6px 0;
            page-break-inside: avoid;
        }}

        /* 分页控制 */
        @media print {{
            body {{
                max-width: 100%;
            }}

            h1, h2, h3, h4, h5, h6 {{
                page-break-after: avoid;
            }}

            table, figure {{
                page-break-inside: avoid;
            }}
        }}

        /* 水平线 */
        hr {{
            border: none;
            border-top: 1px solid #ddd;
            margin: 24px 0;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
        return html_template

    def _generate_pdf_with_weasyprint(self, html_content: str) -> bytes:
        """使用 WeasyPrint 生成 PDF"""
        import weasyprint

        logger.info("🔧 使用 WeasyPrint 生成 PDF...")

        try:
            # 创建 PDF
            pdf_bytes = weasyprint.HTML(string=html_content).write_pdf()

            logger.info(f"✅ WeasyPrint PDF 生成成功，大小: {len(pdf_bytes)} 字节")
            return pdf_bytes
        except OSError as e:
            if "cairo" in str(e).lower():
                error_msg = (
                    "WeasyPrint 缺少 Cairo 库。\n"
                    "Windows 用户请安装 GTK3 运行时:\n"
                    "https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer/releases\n"
                    "或使用其他 PDF 生成工具: pip install pdfkit"
                )
                logger.error(f"❌ {error_msg}")
                raise Exception(error_msg) from e
            else:
                raise

    def _generate_pdf_with_pdfkit(self, html_content: str) -> bytes:
        """使用 pdfkit 生成 PDF"""
        import pdfkit

        logger.info("🔧 使用 pdfkit + wkhtmltopdf 生成 PDF...")

        # 配置选项
        options = {
            'encoding': 'UTF-8',
            'enable-local-file-access': None,
            'page-size': 'A4',
            'margin-top': '20mm',
            'margin-right': '20mm',
            'margin-bottom': '20mm',
            'margin-left': '20mm',
        }

        # 生成 PDF
        pdf_bytes = pdfkit.from_string(html_content, False, options=options)

        logger.info(f"✅ pdfkit PDF 生成成功，大小: {len(pdf_bytes)} 字节")
        return pdf_bytes

    def generate_pdf_report(self, report_doc: Dict[str, Any]) -> bytes:
        """生成 PDF 格式报告（优先使用 WeasyPrint/pdfkit，回退到 Pandoc）"""
        logger.info("📊 开始生成 PDF 文档...")

        # 生成 Markdown 内容
        md_content = self.generate_markdown_report(report_doc)

        # 收集所有错误信息
        errors = []

        # 🔥 策略 1: 优先使用 WeasyPrint（最可靠）
        if self.weasyprint_available:
            try:
                html_content = self._markdown_to_html(md_content)
                return self._generate_pdf_with_weasyprint(html_content)
            except Exception as e:
                error_msg = f"WeasyPrint 生成失败: {e}"
                errors.append(error_msg)
                logger.warning(f"⚠️ {error_msg}，尝试其他方法...")

        # 🔥 策略 2: 使用 pdfkit + wkhtmltopdf
        if self.pdfkit_available:
            try:
                html_content = self._markdown_to_html(md_content)
                return self._generate_pdf_with_pdfkit(html_content)
            except Exception as e:
                error_msg = f"pdfkit 生成失败: {e}"
                errors.append(error_msg)
                logger.warning(f"⚠️ {error_msg}，尝试其他方法...")

        # 🔥 策略 3: 回退到 Pandoc（原有方法）
        if not self.pandoc_available:
            # 生成详细的错误消息
            error_details = "\n".join(errors) if errors else "无可用的 PDF 生成工具"

            help_msg = (
                "所有 PDF 生成工具都不可用。\n\n"
                "推荐解决方案:\n"
                "1. 安装 WeasyPrint (推荐):\n"
                "   pip install weasyprint\n"
                "   Windows 用户还需安装 GTK3 运行时:\n"
                "   https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer/releases\n\n"
                "2. 或安装 pdfkit:\n"
                "   pip install pdfkit\n"
                "   并安装 wkhtmltopdf: https://wkhtmltopdf.org/downloads.html\n\n"
                "3. 或安装 Pandoc:\n"
                "   https://pandoc.org/installing.html\n\n"
                f"错误详情:\n{error_details}"
            )

            logger.error(f"❌ {help_msg}")
            raise Exception(help_msg)

        logger.info("📝 使用 Pandoc 生成 PDF（回退方案）...")

        # PDF 引擎列表（按优先级）
        pdf_engines = [
            ('wkhtmltopdf', 'HTML 转 PDF 引擎（推荐）'),
            ('weasyprint', '现代 HTML 转 PDF 引擎'),
            (None, 'Pandoc 默认引擎')
        ]

        last_error = None

        for engine, description in pdf_engines:
            try:
                # 创建临时文件
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                    output_file = tmp_file.name

                # 创建临时 CSS 文件
                with tempfile.NamedTemporaryFile(mode='w', suffix='.css', delete=False, encoding='utf-8') as css_file:
                    css_file_path = css_file.name
                    css_file.write(self._create_pdf_css())

                # Pandoc 参数
                extra_args = [
                    '--from=markdown-yaml_metadata_block',  # 禁用 YAML 元数据块解析
                    '-V', 'mainfont=Noto Sans CJK SC',  # 设置中文字体（wkhtmltopdf）
                    '-V', 'sansfont=Noto Sans CJK SC',
                    '-V', 'monofont=Noto Sans Mono CJK SC',
                    '--wrap=preserve',  # 保留换行
                    '--columns=120',  # 设置列宽
                    '-V', 'geometry:margin=2cm',  # 设置页边距
                    '-M', 'lang=zh-CN',  # 🔥 明确指定语言为简体中文
                    '-M', 'dir=ltr',  # 🔥 明确指定文本方向为从左到右
                    f'--css={css_file_path}',  # 使用自定义 CSS 控制分页
                ]

                if engine:
                    extra_args.append(f'--pdf-engine={engine}')
                    logger.info(f"🔧 使用 PDF 引擎: {engine}")
                else:
                    logger.info(f"🔧 使用默认 PDF 引擎")

                # 清理内容
                cleaned_content = self._clean_markdown_for_pandoc(md_content)
                
                # 转换为 PDF
                pypandoc.convert_text(
                    cleaned_content,
                    'pdf',
                    format='markdown',
                    outputfile=output_file,
                    extra_args=extra_args
                )
                
                # 检查文件是否生成
                if os.path.exists(output_file) and os.path.getsize(output_file) > 0:
                    # 读取生成的文件
                    with open(output_file, 'rb') as f:
                        pdf_content = f.read()
                    
                    logger.info(f"✅ PDF 生成成功，使用引擎: {engine or '默认'}，大小: {len(pdf_content)} 字节")

                    # 清理临时文件
                    os.unlink(output_file)
                    if 'css_file_path' in locals() and os.path.exists(css_file_path):
                        os.unlink(css_file_path)

                    return pdf_content
                else:
                    raise Exception("PDF 文件生成失败或为空")

            except Exception as e:
                last_error = str(e)
                logger.warning(f"⚠️ PDF 引擎 {engine or '默认'} 失败: {e}")

                # 清理临时文件
                try:
                    if 'output_file' in locals() and os.path.exists(output_file):
                        os.unlink(output_file)
                    if 'css_file_path' in locals() and os.path.exists(css_file_path):
                        os.unlink(css_file_path)
                except:
                    pass

                continue
        
        # 所有引擎都失败
        error_msg = f"""PDF 生成失败，最后错误: {last_error}

可能的解决方案:
1. 安装 wkhtmltopdf (推荐):
   Windows: choco install wkhtmltopdf
   macOS: brew install wkhtmltopdf  
   Linux: sudo apt-get install wkhtmltopdf

2. 安装 LaTeX:
   Windows: choco install miktex
   macOS: brew install mactex
   Linux: sudo apt-get install texlive-full

3. 使用替代格式:
   - Markdown 格式 - 轻量级，兼容性好
   - Word 格式 - 适合进一步编辑
"""
        logger.error(error_msg)
        raise Exception(error_msg)


# 创建全局导出器实例
report_exporter = ReportExporter()

